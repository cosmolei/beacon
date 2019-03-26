import urllib.request
from selenium import webdriver
import os
import pytesser3
import sys, time
from PIL import Image, ImageEnhance
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support.select import Select
from selenium.webdriver.firefox.firefox_binary import FirefoxBinary
import configparser
import json
import docx
import random
import re


# 获取题库数据
def getTkList():
    config = configparser.ConfigParser()
    config.read("settings.conf")
    ## 获取题库列表，选择当前列表中的第一条数据(时间最新)作为当前考试默认题库
    # tkurl = "http://xxjs.dtdjzx.gov.cn/quiz-api/chapter_info/list"
    tkurl = config.get("default", "tk_url")
    tkpage = urllib.request.urlopen(tkurl)
    tkhtml = tkpage.read()
    tkhtmlStr = bytes.decode(tkhtml)
    # print(tkhtmlStr)
    tkjson = json.loads(tkhtmlStr)
    tkid = tkjson['data'][0]['id']

    ##这里加判断，如果tkid为空，初版不做判断，继续执行
    tkdetail_url = config.get("default", "tkdetail_url")
    # tkdetailurl = "http://xxjs.dtdjzx.gov.cn/quiz-api/subject_info/list?chapterId=" + tkid
    tkdetailurl = tkdetail_url + "?chapterId=" + tkid
    tkdetailpage = urllib.request.urlopen(tkdetailurl)
    tkdetailStr = bytes.decode(tkdetailpage.read())
    tkdetailjson = json.loads(tkdetailStr)
    tkdetailList = tkdetailjson['data']['subjectInfoList']

    return tkdetailList

def getWordTkList(tkName, tkList):
    print("开始读取题库 :" + tkName + ".......")
    file = docx.Document(tkName)
    print("段落数：" + str(len(file.paragraphs)))
    fileList = []
    for i in range(len(file.paragraphs)):
        if len(file.paragraphs[i].text) != 0:
            fileList.append(file.paragraphs[i].text)
    pa = 6
    length = len(fileList) // pa

    for i in range(length):
        subjectTitle = fileList[i * pa + 1]
        answer = fileList[i * pa + 6]
        A = fileList[i * pa + 2]
        Aisright = '0'
        if answer.find('A') >= 0:
            Aisright = '1'
        B = fileList[i * pa + 3]
        Bisright = '0'
        if answer.find('B') >= 0:
            Bisright = '1'
        C = fileList[i * pa + 4]
        Cisright = '0'
        if answer.find('C') >= 0:
            Cisright = '1'
        D = fileList[i * pa + 5]
        Disright = '0'
        if answer.find('D') >= 0:
            Disright = '1'
        optionInfoList = [{'optionTitle': A, 'optionType': 'A', 'isRight': Aisright},
                          {'optionTitle': B, 'optionType': 'B', 'isRight': Bisright},
                          {'optionTitle': C, 'optionType': 'C', 'isRight': Cisright},
                          {'optionTitle': D, 'optionType': 'D', 'isRight': Disright}
                          ]
        tkquestion = {'subjectTitle': subjectTitle,
                      'optionInfoList': optionInfoList,
                      'answer': answer, 'A': A, 'B': B, 'C': C, 'D': D}
        tkList.append(tkquestion)
    ##print(tkList)
    print("题库 " + tkName + " 读取完毕.")

def getTkListNew():
    fList = os.listdir()
    tkName = 'new'
    tkList = []
    for i in fList:
        if os.path.splitext(i)[1] == '.docx' or os.path.splitext(i)[1] == '.doc':
            print(i)
            tkName = i
            getWordTkList(tkName, tkList)
    return tkList

def getWordTkListNewNew(tkName, tkList):
    print("开始读取题库 :" + tkName + ".......")
    file = docx.Document(tkName)
    print("段落数：" + str(len(file.paragraphs)))
    fileList = []
    for i in range(len(file.paragraphs)):
        if len(file.paragraphs[i].text) != 0:
            fileList.append(file.paragraphs[i].text)
    pa = 6
    length = len(fileList) // pa
    for i in range(length):
        subjectTitle = fileList[i * pa + 1]
        answer = fileList[i * pa + 6]
        answerString = ""
        A = fileList[i * pa + 2]
        if A.find("、") > 0 and A.find("、") < 3:
            A = A.split("、", 1)[1]
        if A.find(".") > 0 and A.find(".") < 3:
            A = A.split(".", 1)[1]
        Aisright = '0'
        if answer.find('A') >= 0:
            Aisright = '1'
            if answerString == "":
                answerString = answerString + A
            else:
                answerString = answerString + "、" + A
        B = fileList[i * pa + 3]
        if B.find("、") > 0 and B.find("、") < 3:
            B = B.split("、", 1)[1]
        if B.find(".") > 0 and B.find(".") < 3:
            B = B.split(".", 1)[1]
        Bisright = '0'
        if answer.find('B') >= 0:
            Bisright = '1'
            if answerString == "":
                answerString = answerString + B
            else:
                answerString = answerString + "、" + B
        C = fileList[i * pa + 4]
        if C.find("、") > 0 and C.find("、") < 3:
            C = C.split("、", 1)[1]
        if C.find(".") > 0 and C.find(".") < 3:
            C = C.split(".", 1)[1]
        Cisright = '0'
        if answer.find('C') >= 0:
            Cisright = '1'
            if answerString == "":
                answerString = answerString + C
            else:
                answerString = answerString + "、" + C
        D = fileList[i * pa + 5]
        if D.find("、") > 0 and D.find("、") < 3:
            D = D.split("、", 1)[1]
        if D.find(".") > 0 and D.find(".") < 3:
            D = D.split(".", 1)[1]
        Disright = '0'
        if answer.find('D') >= 0:
            Disright = '1'
            if answerString == "":
                answerString = answerString + D
            else:
                answerString = answerString + "、" + D
        optionInfoList = [{'optionTitle': A, 'optionType': 'A', 'isRight': Aisright},
                          {'optionTitle': B, 'optionType': 'B', 'isRight': Bisright},
                          {'optionTitle': C, 'optionType': 'C', 'isRight': Cisright},
                          {'optionTitle': D, 'optionType': 'D', 'isRight': Disright}
                          ]
        subjectDetail = subjectTitle.replace(' ', '').replace('（）', answerString)
        tkquestion = {'subjectTitle': subjectTitle,
                      'subjectDetail': subjectDetail,
                      'optionInfoList': optionInfoList,
                      'answer': answer, 'A': A, 'B': B, 'C': C, 'D': D}
        tkList.append(tkquestion)
    ##print(tkList)
    print("题库 " + tkName + " 读取完毕.")

def getTkListNewNew():
    fList = os.listdir()
    tkName = 'new'
    tkList = []
    for i in fList:
        if os.path.splitext(i)[1] == '.docx' or os.path.splitext(i)[1] == '.doc':
            print(i)
            tkName = i
            getWordTkListNewNew(tkName, tkList)
    return tkList

# 获取浏览器
def getBrowser():
    config = configparser.ConfigParser()
    config.read("settings.conf")
    browserType = config.get("config", "browser")
    if browserType == "Firefox":
        binary = FirefoxBinary('Mozilla Firefox/firefox.exe')
        browser = webdriver.Firefox(firefox_binary=binary, executable_path="geckodriver")
        return browser
    else:
        return None


# 获取问题答案字典
def queryQuestionTitle(title, questionList):
    result_dict = {'optionInfoList': []}
    for i in range(len(questionList)):
        question_dict = questionList[i]
        # if question_dict['subjectTitle'].find(title) >= 0:
        if compareRightToLeft(question_dict['subjectTitle'], title):
            result_dict = question_dict
            break
    return result_dict['optionInfoList']

# 获取问题答案字典new
def queryQuestionDict(dict, questionList):
    title = dict['title']
    list = dict['list']
    result_dict = {}
    for i in range(len(questionList)):
        question_dict = questionList[i]
        # if question_dict['subjectTitle'].find(title) >= 0:
        titles = title.split("（）")
        if len(titles) > 1 :
            if (question_dict['subjectDetail'].find(titles[0]) >= 0) and (question_dict['subjectDetail'].find(titles[1]) >=0) :
                result_dict = question_dict
                break
        # if compareRightToLeft(question_dict['subjectTitle'], title):
        #     result_dict = question_dict
        #     break
        # else:
        #     count = 0
        #     for j in range(len(list)):
        #         if question_dict['subjectDetail'].find(list[j]) >=0:
        #             count = count + 1
        #     if count > 0:
        #         result_dict = question_dict
        #         break
    return result_dict

def isElementExistByCss(r_browser, cssStr):
    flag = True
    browser = r_browser
    try:
        browser.find_element_by_css_selector(cssStr)
        return flag
    except:
        flag = False
        return flag


def isAlertedFuc(r_browser):
    browser = r_browser
    try:
        browser.switch_to_alert()
        return True
    except:
        return False

def randomAnswer(r_browser, count):
    browser = r_browser
    elem_j = random.randint(0, 3)
    option_element = browser.find_element_by_css_selector(
        ".W_ti_ul > li:nth-child(" + str(count) + ") > div:nth-child(" + str(
            elem_j + 2) + ") > label:nth-child(1) > sapn:nth-child(2)")
    option_element.click()

#比较字符串匹配函数
def compareRightToLeft(left, right):
    #定义比较次数
    count_total = 50
    #定义分段数量
    co_ception = 10
    #定义匹配成功率
    count_percent = 0.8
    #定义实际成功次数
    count = 0
    #定义基准分段长度
    str_len_base = 6
    right_value = removeUselessChar(right)
    left_value = removeUselessChar(left)
    for i in range(count_total):
        right_ception_len = len(right_value)//co_ception
        compare_value = right_value[random.randint(0,len(right_value)-right_ception_len):right_ception_len + random.randint(0,len(right_value)-right_ception_len)]
        if left_value.find(compare_value)>=0:
            count = count + 1
    if count > count_percent * count_total:
        return True
    return False

#去除标点符号
def removeUselessChar(value_str):
    temp = value_str
    result_str = re.sub("[\s+\.\!\/_,$%^*(+\"\']+|[+——！，。？、~@#￥%……&*（）]+", "",temp)
    return result_str

# 自动答题主逻辑
def autoExam(r_browser, answer_long, overdure, tkdetailList):
    browser = r_browser
    browser.set_window_size(1280, 1024)  # 防止之前太小遮掩答题选项会造成bug
    count = 1
    no_question = 0
    while count <= 20:
        try:
            if count != 1:
                current_navigate = browser.find_element_by_css_selector(".W_active")
                current_count = int(current_navigate.get_attribute('innerHTML'))
                if current_count > 0 and current_count > count:
                    count = current_count
                    no_question = 0

            if no_question == 0:
                # 导航按键
                elem_navigate = browser.find_element_by_css_selector(".W_kuan > li:nth-child(" + str(count) + ")")
                elem_navigate.click()


            # wait1 = WebDriverWait(browser, overdure).until(EC.presence_of_all_elements_located((By.CSS_SELECTOR, ".W_ti_ul > li:nth-child("+str(count)+") > h1:nth-child(1) > span:nth-child(3)")))
            # wait2 = WebDriverWait(browser, overdure).until(EC.visibility_of_element_located((By.CSS_SELECTOR, ".W_ti_ul > li:nth-child("+str(count)+") > h1:nth-child(1) > span:nth-child(3)")))
            # 查找当前显示的题目
            elem_current_question = browser.find_element_by_css_selector(
                ".W_ti_ul > li:nth-child(" + str(count) + ") > h1:nth-child(1) > span:nth-child(3)")
            # elem_current_question = WebDriverWait(browser, overdure).until(EC.visibility_of_element_located((By.CSS_SELECTOR, ".W_ti_ul > li:nth-child("+str(count)+") > h1:nth-child(1) > span:nth-child(3)")))
            current_question_title = elem_current_question.get_attribute('innerHTML')

            ## current_option_first_read
            current_option_list = []
            for jj in range(4):
                current_option_element = browser.find_element_by_css_selector(
                    ".W_ti_ul > li:nth-child(" + str(count) + ") > div:nth-child(" + str(
                        jj + 2) + ") > label:nth-child(1) > sapn:nth-child(2)")
                current_option_html = current_option_element.get_attribute('innerHTML')
                current_option_value = current_option_html.split('.')[1]
                current_option_list.append(current_option_value)
            current_question_dict = {"title": current_question_title, "list": current_option_list}
            # option_info_list = queryQuestionTitle(current_question_title, tkdetailList)
            # option_info_list = queryQuestionDict(current_question_dict, tkdetailList)
            option_info_dictionary = queryQuestionDict(current_question_dict, tkdetailList)

            # 这里等待做题时间秒数( 加入随机数,防止固定答题时常检查)
            time.sleep(random.randint(1, answer_long + random.randint(1, 3)))
            if option_info_dictionary:
                option_count = 0
                # 获取真实选项结构并进行作答
                for j in range(4):
                    option_element = browser.find_element_by_css_selector(
                        ".W_ti_ul > li:nth-child(" + str(count) + ") > div:nth-child(" + str(
                            j + 2) + ") > label:nth-child(1) > sapn:nth-child(2)")
                    # option_element = WebDriverWait(browser, overdure).until(EC.visibility_of_element_located(
                    #     (By.CSS_SELECTOR, ".W_ti_ul > li:nth-child("+str(count)+") > div:nth-child("+str(j+2)+") > label:nth-child(1) > sapn:nth-child(2)")))
                    option_html = option_element.get_attribute('innerHTML')
                    option_value = option_html.split('.')[1]
                    for k in range(len(option_info_dictionary['optionInfoList'])):
                        option_info_dict = option_info_dictionary['optionInfoList'][k]
                        if (option_info_dict['optionTitle'].find(option_value) >= 0) and (option_info_dict['isRight'] == '1'):
                        # if (compareRightToLeft(option_info_dict['optionTitle'], option_value)) and (option_info_dict['isRight'] == '1'):
                            option_element.click()  # 点击完成选择
                            option_count = option_count + 1
                            print(str(count) + ":" + "option:" + str(j))

                if option_count == 0:
                    # 找不到答案随机一个
                    # print(str(count) + ":no options.")
                    # print("找不到正确选项，请自己选择答案,并点击下一题")

                    # randomAnswer(browser, count)
                    for jj in range(4):
                        option_element = browser.find_element_by_css_selector(
                            ".W_ti_ul > li:nth-child(" + str(count) + ") > div:nth-child(" + str(
                                jj + 2) + ") > label:nth-child(1) > sapn:nth-child(2)")
                        option_html = option_element.get_attribute('innerHTML')
                        option_value = option_html.split('.')[1]
                        if option_info_dictionary['subjectDetail'].find(option_value) >=0:
                            option_element.click()  # 点击完成选择
                            option_count = option_count + 1
                            print(str(count) + ":" + "option:" + str(j))
                    if option_count == 0:
                        print(str(count) + ":no options.")
                        print("找不到正确选项，请自己选择答案,并点击下一题")
                        no_question = 1
                    else:
                        count = count + 1
                else:
                    count = count + 1
            else:
                #找不到答案随机一个
                print(str(count) + ":no this question.")
                print("题库找不到该题干及相关题干, 请自己选择答案,并点击下一题")
                no_question = 1
                # randomAnswer(browser, count)

        except:
            time.sleep(1)
            isAlerted = isAlertedFuc(browser)
            if isAlerted:
                browser.switch_to_alert().accept()
                count = count - 1
            else:
                browser.refresh()
                count = 1

    try:
        elem_submit = browser.find_element_by_css_selector(".W_fr")
        elem_submit.click()

        jiaojuan_button = browser.find_element_by_css_selector(".jiaojuan")
        jiaojuan_button.click()
    except:
        print("非自然结束,请自行结束答题过程")
        anykey = input("=========提交试卷后请按任意键继续============\n")
        if anykey:
            print("现在继续。。。")

# 退出
def exit():
    anykey = input("===================按任意键退出==================\n")
    if anykey:
        sys.exit()


# 生成目录
def pathExamin(path):
    ppath = path
    isExists = os.path.exists(ppath)
    if not isExists:
        os.makedirs(ppath)


# 生成截图
def screenshot_maker(r_browser, user_name):
    browser = r_browser
    path = "image\\" + time.strftime("%Y-%m-%d", time.localtime())
    pathExamin(path)
    browser.get_screenshot_as_file(
        'image\\' + time.strftime("%Y-%m-%d", time.localtime()) + '\\' + user_name + '_points.png')


def isElementExistById(r_browser, element):
    flag = True
    browser = r_browser
    try:
        browser.find_element_by_id(element)
        return flag
    except:
        flag = False
        return flag


def doTheLogin(r_browser, user_dict):
    browser = r_browser
    elem_user = browser.find_element_by_name('username')
    elem_psw = browser.find_element_by_name('password')
    elem_code = browser.find_element_by_name('validateCode')
    elem_yanzhengma = browser.find_element_by_id('yanzhengma')
    # 截图之前再次固定大小

    # -------------------对验证码进行区域截图，好吧，这方法有点low------------------
    pathExamin("tmp")
    browser.get_screenshot_as_file('tmp\\fullimage.png')  # 比较好理解
    im = Image.open('tmp\\fullimage.png')
    # 图片尺寸重置--为了解决不同显示器切割位置不一样的bug
    resize_im = im.resize((995, 668), Image.ANTIALIAS)
    resize_im.save('tmp\\fullimage_resize.png')
    img = Image.open('tmp\\fullimage_resize.png')
    # 设置要裁剪的区域--这里的区域是限制在995*668这个比例的图片尺寸
    #     box = (682,594,837,631)
    box = (545, 474, 670, 507)
    region = img.crop(box)  # 此时，region是一个新的图像对象。
    # region.show()#显示的话就会被占用，所以要注释掉
    region.save("tmp\\image_code.png")

    im = Image.open("tmp\\image_code.png")
    imgry = im.convert('L')  # 图像加强，二值化
    sharpness = ImageEnhance.Contrast(imgry)  # 对比度增强
    sharp_img = sharpness.enhance(2.0)
    sharp_img.save("tmp\\image_code.png")
    # http://www.cnblogs.com/txw1958/archive/2012/02/21/2361330.html
    # imgry.show()#这是分布测试时候用的，整个程序使用需要注释掉
    # imgry.save("D:\\tmp\\image_code.jpg")
    code = pytesser3.image_file_to_string("tmp\\image_code.png")  # code即为识别出的图片数字str类型
    code = code.replace(' ', '')
    print(code)
    # 打印code观察是否识别正确
    # 这里要判断code的字符，如果解析出来不是4个字符的话，则继续刷新页面重新来过。
    # 留个扣子在这里。
    # elem_user.send_keys(user_dict['login_name'])
    # elem_psw.send_keys(user_dict['password'])
    # 改版====数据结构变化
    elem_user.send_keys(user_dict[1])
    elem_psw.send_keys(user_dict[2])
    elem_code.send_keys(code)

    login_button = browser.find_element_by_class_name('tianze-loginbtn')
    login_button.click()
    # 等1秒钟看登陆结果
    wait = WebDriverWait(browser, 1)
    time.sleep(5)
    flag = isElementExistById(browser, 'validateCodeMessage')
    if flag:
        # reinput(browser)
        elem_validate_msg = browser.find_element_by_id('validateCodeMessage').get_attribute('innerHTML')
        if elem_validate_msg.find("验证码") >= 0:
            browser.refresh()
            return doTheLogin(browser, user_dict)
        else:
            print("无法确定的登陆失败")
            return 0
    else:
        print("成功登陆")
        return 1


def write_log(msg):
    fo = open("log.log", "w")
    fo.write(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()) + ":" + msg + "\n")
    fo.close()


def people_exam(master, logText, peopleList):
    # 第一步，声明相关量
    master.reveal_log(logText, "小程序开始读取相关配置......")

    # 删除日志文件
    if os.path.exists("log.log"):
        os.remove("log.log")
    ##获取题库
    master.reveal_log(logText, "开始读取题库......")
    tkdetailList = getTkListNewNew()
    if len(tkdetailList) > 0:
        master.reveal_log(logText, "题库读取完成!")
    else:
        # print("题库读取失败，请联系开发人员!\n")
        master.reveal_log(logText, "题库读取失败，请联系开发人员!")

    # 读取settings
    config = configparser.ConfigParser()
    config.read("settings.conf")
    lightTower_url = config.get("default", "lightTower_url")
    answer_long = config.getint("config", "answer_long")
    overdure = config.getint("config", "overdure")

    # print("正在获取人员信息......\n")
    master.reveal_log(logText, "正在获取人员信息......")
    # file = open('phone.json', "rb")
    # phone_data = json.load(file)
    phone_data = peopleList
    # print("人员信息获取完成，共获得" + str(len(phone_data)) + "条人员信息\n")
    master.reveal_log(logText, "人员信息获取完成，共获得" + str(len(phone_data)) + "条人员信息")
    # print("正在启动答题过程......\n")
    master.reveal_log(logText, "正在启动答题过程......")
    for i in range(len(phone_data)):
        browser = getBrowser()
        browser.implicitly_wait(30)
        browser.set_window_size(1024, 768)  # 设定固定大小
        # 手机号
        # single_phone = phone_data[i]['phone']
        single_phone = peopleList[i]
        # print("正在为手机号为[" + single_phone + "]的人员做答题准备......")
        master.reveal_log(logText, "正在为手机号为[" + single_phone + "]的人员做答题准备......")
        # 网页打开准备工作
        browser.get(lightTower_url)
        time.sleep(5)

        elem_notice_button = browser.find_element_by_css_selector("#l_mods > div:nth-child(1) > button:nth-child(1)")
        elem_notice_button.click()

        time.sleep(5)

        elem_formal_button = browser.find_element_by_id('lbuts')
        elem_formal_button.click()
        # wait = WebDriverWait(browser, 5)
        elem_select_shenfen = browser.find_element_by_id('shenfen')
        # 群众选定
        Select(elem_select_shenfen).select_by_value("1")
        elem_confirm_button = browser.find_element_by_id('bts')
        elem_confirm_button.click()
        # 群众输入手机号
        elem_phone = browser.find_element_by_id('lshouji')
        elem_phone.send_keys(single_phone)

        elem_cf_button = browser.find_element_by_id('btu')
        elem_cf_button.click()

        elem_jihui = browser.find_element_by_css_selector(".l_jihui span")
        jihui_value = int(elem_jihui.get_attribute('innerHTML'))
        if jihui_value > 0:
            # print("人员[" + single_phone + "]可以答题,开始答题")
            master.reveal_log(logText, "人员[" + single_phone + "]可以答题,开始答题")
            elem_formal_button.click()
            autoExam(browser, answer_long, overdure, tkdetailList)
            # 考试成绩截图
            # wait = WebDriverWait(browser, 5)
            time.sleep(5)
            screenshot_maker(browser, single_phone)
            # print("人员[" + single_phone + "]答题完成!")
            master.reveal_log(logText, "人员[" + single_phone + "]答题完成!")
            write_log("人员[" + single_phone + "]答题完成!")
        else:
            screenshot_maker(browser, single_phone)
            fo = open("log.log", "w")
            # print("人员[" + single_phone + "]今天已没有答题机会")
            master.reveal_log(logText, "人员[" + single_phone + "]今天已没有答题机会")
            write_log("人员[" + single_phone + "]今天已没有答题机会")
        # 关闭浏览器
        browser.quit()
    # print("答题结束!\n")
    master.reveal_log(logText, "答题结束!")

def party_exam(master, logText, partyPeopleList):

    master.reveal_log(logText, "小程序开始读取相关配置......")
    # 删除日志文件
    if os.path.exists("log.log"):
        os.remove("log.log")
    ##获取题库
    master.reveal_log(logText, "开始读取题库......")
    tkdetailList = getTkListNewNew()
    if len(tkdetailList) > 0:
        master.reveal_log(logText, "题库读取完成!")
    else:
        master.reveal_log(logText, "题库读取失败，请联系开发人员!")

    # 读取settings
    config = configparser.ConfigParser()
    config.read("settings.conf")
    login_url = config.get("default", "login_url")
    lightTower_url = config.get("default", "lightTower_url")
    logout_url = config.get("default", "logout_url")
    answer_long = config.getint("config", "answer_long")
    overdure = config.getint("config", "overdure")
    wrong_time = config.getint("config", "wrong_time")

    master.reveal_log(logText, "正在获取人员信息......")
    # file = open('user.json', "rb")
    # user_data = json.load(file)
    user_data = partyPeopleList
    master.reveal_log(logText, "人员信息获取完成，共获得" + str(len(user_data)) + "条人员信息\n")
    master.reveal_log(logText, "正在启动答题过程......")

    for i in range(len(user_data)):
        user_name = user_data[i][0]
        browser = getBrowser()
        browser.set_window_size(1024, 768)  # 设定固定大小

        # 准备登陆
        browser.get(login_url)
        isLogined = doTheLogin(browser, user_data[i])
        if isLogined == 0:
            master.reveal_log(logText, "人员[" + user_name + "]尝试登陆失败，退出本次答题列表！")
            continue

        # 登陆成功继续答题
        master.reveal_log(logText, "人员[" + user_name + "]登陆成功！准备进行答题")
        browser.get(lightTower_url)

        time.sleep(5)

        elem_notice_button = browser.find_element_by_css_selector("#l_mods > div:nth-child(1) > button:nth-child(1)")
        elem_notice_button.click()

        time.sleep(5)

        elem_formal_button = browser.find_element_by_id('lbuts')
        elem_formal_button.click()
        elem_select_shenfen = browser.find_element_by_id('shenfen')
        # 党员选定
        Select(elem_select_shenfen).select_by_value("0")

        elem_confirm_button = browser.find_element_by_id('bts')
        elem_confirm_button.click()

        browser.get(lightTower_url)
        time.sleep(5)

        elem_notice_button = browser.find_element_by_css_selector("#l_mods > div:nth-child(1) > button:nth-child(1)")
        elem_notice_button.click()

        time.sleep(5)

        elem_jihui = browser.find_element_by_css_selector(".l_jihui span")
        jihui_value = int(elem_jihui.get_attribute('innerHTML'))
        if jihui_value > 0:
            master.reveal_log(logText, "人员[" + user_name + "]可以答题,跳转答题")
            time.sleep(5)
            elem_formal2_button = browser.find_element_by_id('lbuts')
            elem_formal2_button.click()
            # 考试主要过程
            autoExam(browser, answer_long, overdure, tkdetailList)
            time.sleep(5)
            screenshot_maker(browser, user_name)
            master.reveal_log(logText, "人员[" + user_name + "]答题完成!")
            write_log("人员[" + user_name + "]答题完成!")
        else:
            master.reveal_log(logText, "人员[" + user_name + "]今天已没有答题机会!")
            write_log("人员[" + user_name + "]今天已没有答题机会!")
            # 准备退出当前账号
        browser.get(logout_url)
        browser.quit()
    master.reveal_log(logText, "自动答题结束!")